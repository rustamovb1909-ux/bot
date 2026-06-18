import os
import sqlite3
import asyncio
import datetime
import tempfile
import shutil
import random
import re
from pathlib import Path
from threading import Thread

from flask import Flask, request, jsonify
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, filters, ContextTypes
)
from telegram.constants import ParseMode

# DOCX uchun
try:
    from docx import Document
except ImportError:
    Document = None

# ==================== KONFIGURATSIYA ====================
TOKEN = os.getenv("BOT_TOKEN") or os.getenv("TOKEN")
if not TOKEN:
    raise ValueError("BOT_TOKEN muhit o'zgaruvchisi o'rnatilmagan!")

WEBHOOK_URL = os.getenv("WEBHOOK_URL", "https://bot-gujm.onrender.com")
PORT = int(os.environ.get("PORT", 10000))

app = Flask(__name__)

# ==================== DATABASE ====================
DB_PATH = "data/test_bot.db"
os.makedirs("data", exist_ok=True)
os.makedirs("uploads", exist_ok=True)

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        user_id INTEGER PRIMARY KEY,
        username TEXT,
        first_name TEXT,
        last_name TEXT,
        registered_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        file_name TEXT,
        file_path TEXT,
        file_size INTEGER,
        original_name TEXT,
        format TEXT,
        question_count INTEGER DEFAULT 0,
        uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(user_id)
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS test_questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        file_id INTEGER,
        question_text TEXT,
        option_a TEXT,
        option_b TEXT,
        option_c TEXT,
        option_d TEXT,
        correct_answer TEXT,
        FOREIGN KEY (file_id) REFERENCES files(id) ON DELETE CASCADE
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS test_sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        file_id INTEGER,
        total_questions INTEGER,
        correct_answers INTEGER DEFAULT 0,
        wrong_answers INTEGER DEFAULT 0,
        skipped_answers INTEGER DEFAULT 0,
        score REAL DEFAULT 0,
        status TEXT DEFAULT 'in_progress',
        started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        finished_at TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(user_id),
        FOREIGN KEY (file_id) REFERENCES files(id)
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS test_answers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        session_id INTEGER,
        question_number INTEGER,
        question_text TEXT,
        user_answer TEXT,
        correct_answer TEXT,
        is_correct BOOLEAN,
        FOREIGN KEY (session_id) REFERENCES test_sessions(id)
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS user_settings (
        user_id INTEGER PRIMARY KEY,
        default_test_count INTEGER DEFAULT 10,
        shuffle_questions BOOLEAN DEFAULT 1,
        shuffle_options BOOLEAN DEFAULT 1,
        FOREIGN KEY (user_id) REFERENCES users(user_id)
    )''')
    
    conn.commit()
    conn.close()

init_db()

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

# ==================== PARSERLAR ====================

def parse_txt_file(file_path):
    questions = []
    
    content = None
    for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, LookupError):
            pass
    
    if not content:
        return []
    
    lines = content.splitlines()
    non_empty = [l.strip() for l in lines if l.strip()]
    
    if not non_empty:
        return []
    
    has_hash = any(l.startswith("#") for l in non_empty)
    has_plus = any(l.startswith("+") for l in non_empty)
    has_pipe = any("|" in l and l.count("|") >= 4 for l in non_empty)
    has_numbered = any(re.match(r"^\d+[.)]\s+", l) for l in non_empty)
    has_abcd = any(re.match(r"^[A-Da-d][.)]\s+", l) for l in non_empty)
    
    if has_hash and has_plus:
        questions = parse_hash_format(lines)
        if questions:
            return questions
    
    if has_numbered and has_abcd:
        questions = parse_numbered_format(lines)
        if questions:
            return questions
    
    if has_pipe:
        questions = parse_pipe_format(non_empty)
        if questions:
            return questions
    
    questions = parse_qa_format(lines)
    if questions:
        return questions
    
    for parser in [parse_hash_format, parse_numbered_format, parse_qa_format]:
        questions = parser(lines)
        if questions:
            return questions
    
    questions = parse_pipe_format(non_empty)
    return questions


def parse_hash_format(lines):
    questions = []
    current_q = None
    correct = None
    opts = []
    state = "idle"
    
    def flush():
        nonlocal current_q, correct, opts, state
        if current_q and correct and len(opts) >= 2:
            all_opts = opts[:4]
            while len(all_opts) < 4:
                all_opts.append(f"Variant {len(all_opts) + 1}")
            if correct not in all_opts:
                all_opts.insert(0, correct)
                all_opts = all_opts[:4]
            questions.append({
                "text": current_q,
                "options": all_opts,
                "correct": correct
            })
        current_q, correct, opts, state = None, None, [], "idle"
    
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        
        if line.startswith("#") or line.startswith("?"):
            flush()
            q = line[1:].strip()
            q = re.sub(r"\s*\?$", "", q).strip()
            if q:
                current_q = q
                state = "idle"
            else:
                state = "need_q"
        
        elif state == "need_q":
            current_q = re.sub(r"\s*\?$", "", line).strip()
            state = "idle"
        
        elif line.startswith("+"):
            ans = line[1:].strip()
            if not ans:
                state = "need_correct"
            elif current_q is not None:
                correct = ans
                if ans not in opts:
                    opts.append(ans)
                state = "idle"
        
        elif state == "need_correct":
            correct = line
            if line not in opts:
                opts.append(line)
            state = "idle"
        
        elif line.startswith("-"):
            ans = line[1:].strip()
            if not ans:
                state = "need_wrong"
            elif current_q is not None and ans not in opts:
                opts.append(ans)
        
        elif state == "need_wrong":
            if current_q is not None and line not in opts:
                opts.append(line)
            state = "idle"
    
    flush()
    return questions


def parse_numbered_format(lines):
    questions = []
    current_q = None
    opts_dict = {}
    correct_letter = None
    
    def flush():
        nonlocal current_q, opts_dict, correct_letter
        if current_q and opts_dict and correct_letter:
            ul = correct_letter.upper()
            if ul in opts_dict:
                correct_ans = opts_dict[ul]
                options = list(opts_dict.values())[:4]
                while len(options) < 4:
                    options.append(f"Variant {len(options) + 1}")
                questions.append({
                    "text": current_q,
                    "options": options,
                    "correct": correct_ans
                })
        current_q, opts_dict, correct_letter = None, {}, None
    
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        
        m = re.match(r"^(\d+)[.)]\s+(.+)$", line)
        if m:
            flush()
            current_q = m.group(2).strip()
            continue
        
        m = re.match(r"^([A-Da-d])[.)]\s+(.+)$", line)
        if m:
            opts_dict[m.group(1).upper()] = m.group(2).strip()
            continue
        
        m = re.match(
            r"^(?:Javob|To'g'ri\s*javob|Answer|Ans|Togri\s*javob|Javobi)[:\s]*([A-Da-d])",
            line, re.IGNORECASE
        )
        if m:
            correct_letter = m.group(1).upper()
    
    flush()
    return questions


def parse_pipe_format(lines):
    questions = []
    for line in lines:
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 5 and parts[0]:
            questions.append({
                "text": parts[0],
                "options": parts[1:5],
                "correct": parts[1]
            })
    return questions


def parse_qa_format(lines):
    questions = []
    current_q = None
    current_ans = None
    
    def flush():
        nonlocal current_q, current_ans
        if current_q and current_ans:
            questions.append({
                "text": current_q,
                "options": [current_ans, "Variant B", "Variant C", "Variant D"],
                "correct": current_ans
            })
        current_q, current_ans = None, None
    
    for line in lines:
        line = line.strip()
        if not line:
            flush()
            continue
        
        if re.match(r"^(savol|question)\s*:", line, re.IGNORECASE):
            flush()
            current_q = re.split(r":\s*", line, maxsplit=1)[-1].strip()
        elif re.match(r"^(javob|answer)\s*:", line, re.IGNORECASE):
            current_ans = re.split(r":\s*", line, maxsplit=1)[-1].strip()
    
    flush()
    return questions


def parse_docx_file(file_path):
    if not Document:
        return []
    
    try:
        doc = Document(file_path)
        questions = []
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                cells = [c for c in cells if c]
                if len(cells) >= 5:
                    questions.append({
                        "text": cells[0],
                        "options": cells[1:5],
                        "correct": cells[1]
                    })
        
        if not questions:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for parser in [parse_hash_format, parse_numbered_format, parse_qa_format]:
                questions = parser(lines)
                if questions:
                    break
        
        return questions
    except Exception as e:
        print(f"DOCX o'qishda xatolik: {e}")
        return []


# ==================== BOT HANDLERLAR ====================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT OR IGNORE INTO users (user_id, username, first_name, last_name)
                 VALUES (?, ?, ?, ?)''',
              (user.id, user.username, user.first_name, user.last_name))
    c.execute('''INSERT OR IGNORE INTO user_settings (user_id) VALUES (?)''', (user.id,))
    conn.commit()
    conn.close()
    
    keyboard = [
        [InlineKeyboardButton("📤 Test fayl yuklash", callback_data="upload_file")],
        [InlineKeyboardButton("📊 Mening testlarim", callback_data="my_tests")],
        [InlineKeyboardButton("📈 Natijalarim", callback_data="my_results")],
        [InlineKeyboardButton("📁 Fayllarim", callback_data="my_files")],
        [InlineKeyboardButton("⚙️ Sozlamalar", callback_data="settings")],
        [InlineKeyboardButton("💡 Yordam", callback_data="help")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"👋 Assalomu alaykum, {user.first_name}!\n\n"
        "🎯 <b>TEST MASTER BOT</b> ga xush kelibsiz!\n\n"
        "✨ <b>Imkoniyatlar:</b>\n"
        "• 📄 Test fayllarni yuklash (TXT, DOCX)\n"
        "• 🎲 Savollar va variantlar aralash\n"
        "• 📝 Test sonini o'zingiz belgilaysiz\n"
        "• 📊 Batafsil statistika\n"
        "• 📁 Fayllar tarixi\n\n"
        "📎 <b>Qo'llab-quvvatlanadigan formatlar:</b>\n"
        "• <code>.txt</code> — Matnli fayl\n"
        "• <code>.docx</code> — Word\n\n"
        "Boshlash uchun menyudan tanlang 👇",
        parse_mode=ParseMode.HTML,
        reply_markup=reply_markup
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📖 <b>YORDAM</b>\n\n"
        "<b>📝 TXT Format 1</b> (# belgisi):\n"
        "<code># Savol matni\n+ To'g'ri javob\n- Noto'g'ri 1\n- Noto'g'ri 2\n- Noto'g'ri 3</code>\n\n"
        "<b>📝 TXT Format 2</b> (A/B/C/D):\n"
        "<code>1. Savol matni\nA) To'g'ri\nB) Noto'g'ri 1\nC) Noto'g'ri 2\nD) Noto'g'ri 3\nJavob: A</code>\n\n"
        "<b>📝 TXT Format 3</b> (Pipe):\n"
        "<code>Savol|A variant|B variant|C variant|D variant</code>\n\n"
        "<b>📘 DOCX:</b> 5 ustunli jadval\n(1-ustun: Savol, 2-5: Variantlar)",
        parse_mode=ParseMode.HTML
    )


async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    data = query.data
    user_id = query.from_user.id
    
    if data == "upload_file":
        await query.edit_message_text(
            "📤 <b>Test faylini yuklang</b>\n\n"
            "Faylni yuboring (TXT yoki DOCX).\n\n"
            "📝 Format haqida ma'lumot uchun /help",
            parse_mode=ParseMode.HTML
        )
        context.user_data['waiting_for_file'] = True
    
    elif data == "my_tests":
        await show_available_tests(query, user_id)
    
    elif data == "my_results":
        await show_user_results(query, user_id)
    
    elif data == "my_files":
        await show_user_files(query, user_id)
    
    elif data == "settings":
        await show_settings(query, user_id)
    
    elif data == "help":
        await show_help(query)
    
    elif data == "back_to_main":
        await show_main_menu(query)
    
    elif data.startswith("test_"):
        file_id = int(data.split("_")[1])
        await prepare_test(query, context, file_id)
    
    elif data.startswith("set_count_"):
        count = int(data.split("_")[2])
        await set_test_count(query, user_id, count)
    
    elif data.startswith("start_test_"):
        parts = data.split("_")
        file_id = int(parts[2])
        count_param = parts[3] if len(parts) > 3 else "10"
        await start_test(query, context, file_id, count_param)
    
    elif data.startswith("delete_file_"):
        file_id = int(data.split("_")[2])
        await delete_file(query, user_id, file_id)
    
    elif data.startswith("answer_"):
        await handle_answer(query, context, data)
    
    elif data == "skip_question":
        await skip_question(query, context)
    
    elif data == "finish_test":
        await finish_test(query, context)
    
    elif data.startswith("retry_"):
        file_id = int(data.split("_")[1])
        await prepare_test(query, context, file_id)
    
    elif data == "cancel_test":
        await cancel_test(query, context)


async def show_main_menu(query):
    keyboard = [
        [InlineKeyboardButton("📤 Test fayl yuklash", callback_data="upload_file")],
        [InlineKeyboardButton("📊 Mening testlarim", callback_data="my_tests")],
        [InlineKeyboardButton("📈 Natijalarim", callback_data="my_results")],
        [InlineKeyboardButton("📁 Fayllarim", callback_data="my_files")],
        [InlineKeyboardButton("⚙️ Sozlamalar", callback_data="settings")],
        [InlineKeyboardButton("💡 Yordam", callback_data="help")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.edit_message_text(
        "🏠 <b>ASOSIY MENYU</b>\n\nTest tizimiga xush kelibsiz!",
        parse_mode=ParseMode.HTML,
        reply_markup=reply_markup
    )


async def show_help(query):
    keyboard = [[InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")]]
    await query.edit_message_text(
        "📖 <b>YORDAM</b>\n\n"
        "<b>📝 TXT Format 1</b> (# belgisi):\n"
        "<code># Savol matni\n+ To'g'ri javob\n- Noto'g'ri 1\n- Noto'g'ri 2\n- Noto'g'ri 3</code>\n\n"
        "<b>📝 TXT Format 2</b> (A/B/C/D):\n"
        "<code>1. Savol matni\nA) To'g'ri\nB) Noto'g'ri 1\nC) Noto'g'ri 2\nD) Noto'g'ri 3\nJavob: A</code>\n\n"
        "<b>📝 TXT Format 3</b> (Pipe):\n"
        "<code>Savol|A variant|B variant|C variant|D variant</code>\n\n"
        "<b>📘 DOCX:</b> 5 ustunli jadval",
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def show_available_tests(query, user_id):
    conn = get_db()
    c = conn.cursor()
    files = c.execute('''SELECT id, original_name, question_count, uploaded_at
                         FROM files WHERE user_id = ? ORDER BY uploaded_at DESC''',
                      (user_id,)).fetchall()
    conn.close()
    
    if not files:
        keyboard = [[InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")]]
        await query.edit_message_text(
            "📂 Siz hali hech qanday test fayl yuklamagansiz.\n\n"
            "📤 Fayl yuklash uchun 'Test fayl yuklash' tugmasini bosing.",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return
    
    text = "📊 <b>TEST BOSHLASH</b>\n\nFaylni tanlang:\n\n"
    keyboard = []
    
    for i, f in enumerate(files, 1):
        text += f"<b>{i}.</b> {f['original_name']} — <b>{f['question_count']} savol</b>\n"
        keyboard.append([
            InlineKeyboardButton(
                f"▶️ {f['original_name'][:30]}",
                callback_data=f"test_{f['id']}"
            )
        ])
    
    keyboard.append([InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")])
    
    await query.edit_message_text(
        text,
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def show_user_files(query, user_id):
    conn = get_db()
    c = conn.cursor()
    files = c.execute('''SELECT id, original_name, question_count, uploaded_at
                         FROM files WHERE user_id = ? ORDER BY uploaded_at DESC LIMIT 10''',
                      (user_id,)).fetchall()
    conn.close()
    
    if not files:
        keyboard = [[InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")]]
        await query.edit_message_text(
            "📁 Hali hech qanday fayl yuklanmagan.",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return
    
    text = "📁 <b>MENING FAYLLARIM</b>\n\n"
    keyboard = []
    
    for f in files:
        text += f"📄 {f['original_name']} — {f['question_count']} savol\n"
        text += f"   📅 {f['uploaded_at'][:16]}\n\n"
        keyboard.append([
            InlineKeyboardButton(f"🗑️ {f['original_name'][:25]}", callback_data=f"delete_file_{f['id']}")
        ])
    
    keyboard.append([InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")])
    
    await query.edit_message_text(
        text,
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def show_user_results(query, user_id):
    conn = get_db()
    c = conn.cursor()
    results = c.execute('''SELECT ts.*, f.original_name
                          FROM test_sessions ts
                          JOIN files f ON ts.file_id = f.id
                          WHERE ts.user_id = ? AND ts.status = 'completed'
                          ORDER BY ts.finished_at DESC LIMIT 10''',
                       (user_id,)).fetchall()
    
    total_tests = len(results)
    avg_score = sum(r['score'] for r in results) / len(results) if results else 0
    total_correct = sum(r['correct_answers'] for r in results)
    total_questions = sum(r['total_questions'] for r in results)
    
    conn.close()
    
    keyboard = [[InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")]]
    
    if not results:
        await query.edit_message_text(
            "📊 Siz hali hech qanday test topshirmagansiz.",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return
    
    text = "📊 <b>NATIJALARIM</b>\n\n"
    text += f"📝 Jami testlar: <b>{total_tests} ta</b>\n"
    text += f"📈 O'rtacha: <b>{avg_score:.1f}%</b>\n"
    text += f"✅ Jami to'g'ri: <b>{total_correct}</b> / {total_questions}\n\n"
    text += "━" * 30 + "\n\n"
    
    for i, r in enumerate(results[:5], 1):
        text += f"<b>{i}.</b> {r['original_name']}\n"
        text += f"   ✅ {r['correct_answers']} | ❌ {r['wrong_answers']} | ⏭️ {r['skipped_answers']}\n"
        text += f"   📊 {r['score']:.1f}%\n\n"
    
    await query.edit_message_text(
        text,
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def show_settings(query, user_id):
    conn = get_db()
    c = conn.cursor()
    settings = c.execute('SELECT * FROM user_settings WHERE user_id = ?', (user_id,)).fetchone()
    conn.close()
    
    current_count = settings['default_test_count'] if settings else 10
    count_label = "Barchasi" if current_count == 0 else f"{current_count} ta"
    
    keyboard = [
        [InlineKeyboardButton(f"📝 10 ta savol {'✅' if current_count == 10 else ''}", callback_data="set_count_10")],
        [InlineKeyboardButton(f"📝 25 ta savol {'✅' if current_count == 25 else ''}", callback_data="set_count_25")],
        [InlineKeyboardButton(f"📝 50 ta savol {'✅' if current_count == 50 else ''}", callback_data="set_count_50")],
        [InlineKeyboardButton(f"📝 Barchasi {'✅' if current_count == 0 else ''}", callback_data="set_count_0")],
        [InlineKeyboardButton("🔙 Ortga", callback_data="back_to_main")]
    ]
    
    await query.edit_message_text(
        f"⚙️ <b>SOZLAMALAR</b>\n\n"
        f"📝 Testlar soni: <b>{count_label}</b>\n\n"
        "Test boshlaganda nechta savol olishni tanlang:",
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def set_test_count(query, user_id, count):
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT OR REPLACE INTO user_settings (user_id, default_test_count)
                 VALUES (?, ?)''', (user_id, count))
    conn.commit()
    conn.close()
    
    label = "Barchasi" if count == 0 else f"{count} ta"
    await query.answer(f"✅ Testlar soni {label} ga o'rnatildi")
    await show_settings(query, user_id)


async def prepare_test(query, context, file_id):
    user_id = query.from_user.id
    
    conn = get_db()
    c = conn.cursor()
    file_info = c.execute('SELECT * FROM files WHERE id = ? AND user_id = ?', 
                          (file_id, user_id)).fetchone()
    settings = c.execute('SELECT * FROM user_settings WHERE user_id = ?', (user_id,)).fetchone()
    conn.close()
    
    if not file_info:
        await query.answer("❌ Fayl topilmadi!", show_alert=True)
        return
    
    total = file_info['question_count']
    default_count = settings['default_test_count'] if settings else 10
    
    if default_count > total or default_count == 0:
        default_count = total
    
    opts = sorted({n for n in (5, 10, 15, 20, 25, 30, 50) if n <= total} | {total})
    
    keyboard = []
    row = []
    for i, c in enumerate(opts):
        label = f"📚 Hammasi ({c})" if c == total else f"📝 {c} ta"
        row.append(InlineKeyboardButton(label, callback_data=f"start_test_{file_id}_{c}"))
        if len(row) == 2 or i == len(opts) - 1:
            keyboard.append(row)
            row = []
    
    keyboard.append([InlineKeyboardButton("🎲 Tasodifiy", callback_data=f"start_test_{file_id}_random")])
    keyboard.append([InlineKeyboardButton("❌ Bekor qilish", callback_data="my_tests")])
    
    await query.edit_message_text(
        f"📝 <b>TEST SONINI TANLANG</b>\n\n"
        f"📄 Fayl: {file_info['original_name']}\n"
        f"📚 Jami savollar: <b>{total} ta</b>\n\n"
        f"<i>Variantni tanlang 👇</i>",
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def start_test(query, context, file_id, count_param):
    user_id = query.from_user.id
    
    conn = get_db()
    c = conn.cursor()
    
    questions = c.execute('''SELECT * FROM test_questions WHERE file_id = ?''', 
                         (file_id,)).fetchall()
    settings = c.execute('SELECT * FROM user_settings WHERE user_id = ?', (user_id,)).fetchone()
    conn.close()
    
    if not questions:
        await query.edit_message_text("❌ Bu faylda savollar topilmadi.")
        return
    
    questions = list(questions)
    total_available = len(questions)
    
    if count_param == "random":
        count = random.randint(min(5, total_available), min(50, total_available))
    else:
        count = int(count_param)
        if count > total_available:
            count = total_available
    
    if settings and settings['shuffle_questions']:
        random.shuffle(questions)
    
    selected = questions[:count]
    
    processed_questions = []
    for q in selected:
        options = [q['option_a'], q['option_b'], q['option_c'], q['option_d']]
        correct = q['correct_answer']
        
        if settings and settings['shuffle_options']:
            random.shuffle(options)
        
        try:
            correct_idx = options.index(correct)
        except ValueError:
            correct_idx = 0
            options[0] = correct
        
        processed_questions.append({
            'id': q['id'],
            'text': q['question_text'],
            'options': options,
            'correct': correct,
            'correct_idx': correct_idx
        })
    
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO test_sessions (user_id, file_id, total_questions)
                 VALUES (?, ?, ?)''', (user_id, file_id, count))
    session_id = c.lastrowid
    conn.commit()
    conn.close()
    
    context.user_data['test_session_id'] = session_id
    context.user_data['test_file_id'] = file_id
    context.user_data['test_questions'] = processed_questions
    context.user_data['test_current'] = 0
    context.user_data['test_correct'] = 0
    context.user_data['test_wrong'] = 0
    context.user_data['test_skipped'] = 0
    context.user_data['test_answers'] = []
    
    await query.edit_message_text(
        f"🚀 <b>TEST BOSHLANDI!</b>\n\n"
        f"📝 Jami: <b>{count} ta savol</b>\n"
        f"🎲 Savollar aralash holda\n\n"
        f"<i>Omad! 🍀</i>",
        parse_mode=ParseMode.HTML
    )
    
    await asyncio.sleep(1.5)
    await show_question(query, context)


async def show_question(query, context):
    questions = context.user_data.get('test_questions', [])
    current = context.user_data.get('test_current', 0)
    
    if current >= len(questions):
        await finish_test(query, context)
        return
    
    q = questions[current]
    
    text = f"📝 <b>Savol {current + 1}/{len(questions)}</b>\n\n"
    text += f"<b>{q['text']}</b>\n\n"
    
    keyboard = []
    for i, opt in enumerate(q['options']):
        letter = chr(65 + i)
        if opt:
            display_opt = opt[:50] + ('...' if len(opt) > 50 else '')
            keyboard.append([InlineKeyboardButton(
                f"{letter}) {display_opt}",
                callback_data=f"answer_{i}_{q['correct_idx']}"
            )])
    
    keyboard.append([
        InlineKeyboardButton("⏭️ O'tkazish", callback_data="skip_question"),
        InlineKeyboardButton("🔴 Yakunlash", callback_data="finish_test")
    ])
    
    await query.edit_message_text(
        text,
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


async def handle_answer(query, context, data):
    parts = data.split("_")
    user_choice = int(parts[1])
    correct_idx = int(parts[2])
    
    questions = context.user_data.get('test_questions', [])
    current = context.user_data.get('test_current', 0)
    
    if current >= len(questions):
        return
    
    q = questions[current]
    is_correct = (user_choice == correct_idx)
    
    if is_correct:
        context.user_data['test_correct'] = context.user_data.get('test_correct', 0) + 1
    else:
        context.user_data['test_wrong'] = context.user_data.get('test_wrong', 0) + 1
    
    context.user_data['test_answers'] = context.user_data.get('test_answers', [])
    context.user_data['test_answers'].append({
        'question_text': q['text'],
        'user_answer': q['options'][user_choice],
        'correct_answer': q['correct'],
        'is_correct': is_correct
    })
    
    context.user_data['test_current'] = current + 1
    
    emoji = "✅" if is_correct else "❌"
    await query.answer(f"{emoji} {'To\'g\'ri!' if is_correct else 'Noto\'g\'ri!'}")
    
    await show_question(query, context)


async def skip_question(query, context):
    questions = context.user_data.get('test_questions', [])
    current = context.user_data.get('test_current', 0)
    
    if current < len(questions):
        q = questions[current]
        context.user_data['test_answers'] = context.user_data.get('test_answers', [])
        context.user_data['test_answers'].append({
            'question_text': q['text'],
            'user_answer': "O'tkazildi",
            'correct_answer': q['correct'],
            'is_correct': False
        })
    
    context.user_data['test_skipped'] = context.user_data.get('test_skipped', 0) + 1
    context.user_data['test_current'] = current + 1
    
    await query.answer("⏭️ O'tkazildi")
    await show_question(query, context)


async def finish_test(query, context):
    correct = context.user_data.get('test_correct', 0)
    wrong = context.user_data.get('test_wrong', 0)
    skipped = context.user_data.get('test_skipped', 0)
    total = len(context.user_data.get('test_questions', []))
    session_id = context.user_data.get('test_session_id')
    file_id = context.user_data.get('test_file_id')
    answers = context.user_data.get('test_answers', [])
    
    if total == 0:
        await query.edit_message_text("❌ Xatolik yuz berdi.")
        return
    
    score = (correct / total) * 100
    
    if score >= 90:
        grade, emoji = "A'lo", "🏆"
    elif score >= 75:
        grade, emoji = "Yaxshi", "🎉"
    elif score >= 60:
        grade, emoji = "Qoniqarli", "👍"
    else:
        grade, emoji = "O'qish kerak", "📚"
    
    if session_id:
        conn = get_db()
        c = conn.cursor()
        c.execute('''UPDATE test_sessions SET 
                     correct_answers = ?, wrong_answers = ?, skipped_answers = ?,
                     score = ?, status = 'completed', finished_at = CURRENT_TIMESTAMP
                     WHERE id = ?''',
                  (correct, wrong, skipped, score, session_id))
        
        for i, ans in enumerate(answers):
            c.execute('''INSERT INTO test_answers 
                         (session_id, question_number, question_text, user_answer, correct_answer, is_correct)
                         VALUES (?, ?, ?, ?, ?, ?)''',
                      (session_id, i + 1, ans['question_text'][:200], 
                       ans['user_answer'][:100], ans['correct_answer'][:100], ans['is_correct']))
        
        conn.commit()
        conn.close()
    
    text = f"{emoji} <b>TEST NATIJASI</b>\n\n"
    text += f"📊 <b>Statistika:</b>\n"
    text += f"• Jami: <b>{total} ta</b>\n"
    text += f"• To'g'ri: <b>{correct} ta</b> ✅\n"
    text += f"• Noto'g'ri: <b>{wrong} ta</b> ❌\n"
    text += f"• O'tkazilgan: <b>{skipped} ta</b> ⏭️\n\n"
    text += f"📈 Foiz: <b>{score:.1f}%</b>\n"
    text += f"🏆 Baho: <b>{grade}</b>"
    
    keyboard = [
        [
            InlineKeyboardButton("🔄 Qayta urinish", callback_data=f"retry_{file_id}"),
            InlineKeyboardButton("🏠 Menyu", callback_data="back_to_main")
        ]
    ]
    
    await query.edit_message_text(
        text,
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    for key in ['test_questions', 'test_current', 'test_correct', 'test_wrong', 
                'test_skipped', 'test_answers', 'test_session_id', 'test_file_id']:
        context.user_data.pop(key, None)


async def cancel_test(query, context):
    for key in ['test_questions', 'test_current', 'test_correct', 'test_wrong', 
                'test_skipped', 'test_answers', 'test_session_id', 'test_file_id']:
        context.user_data.pop(key, None)
    
    await query.edit_message_text("❌ Test bekor qilindi.")
    await show_main_menu(query)


async def delete_file(query, user_id, file_id):
    conn = get_db()
    c = conn.cursor()
    file_info = c.execute('SELECT * FROM files WHERE id = ? AND user_id = ?', 
                          (file_id, user_id)).fetchone()
    
    if file_info:
        if file_info['file_path'] and os.path.exists(file_info['file_path']):
            try:
                os.unlink(file_info['file_path'])
            except Exception:
                pass
        
        c.execute('DELETE FROM test_answers WHERE session_id IN (SELECT id FROM test_sessions WHERE file_id = ?)', (file_id,))
        c.execute('DELETE FROM test_sessions WHERE file_id = ?', (file_id,))
        c.execute('DELETE FROM test_questions WHERE file_id = ?', (file_id,))
        c.execute('DELETE FROM files WHERE id = ?', (file_id,))
        conn.commit()
    
    conn.close()
    
    await query.answer("✅ Fayl o'chirildi")
    await show_user_files(query, user_id)


async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.user_data.get('waiting_for_file'):
        return
    
    document = update.message.document
    if not document:
        await update.message.reply_text("❌ Iltimos, fayl yuboring.")
        return
    
    file_name = document.file_name
    file_size = document.file_size
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext not in ['.txt', '.docx']:
        await update.message.reply_text(
            f"❌ {ext} format qo'llab-quvvatlanmaydi.\n"
            "Qo'llab-quvvatlanadigan formatlar: <b>TXT, DOCX</b>",
            parse_mode=ParseMode.HTML
        )
        return
    
    processing_msg = await update.message.reply_text("⏳ Fayl yuklanmoqda va tahlil qilinmoqda...")
    
    try:
        file = await context.bot.get_file(document.file_id)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            await file.download_to_drive(tmp_file.name)
            tmp_path = tmp_file.name
        
        if ext == '.txt':
            questions = parse_txt_file(tmp_path)
        else:
            questions = parse_docx_file(tmp_path)
        
        if not questions:
            await processing_msg.edit_text(
                "❌ Fayldan savollar topilmadi.\n\n"
                "📝 To'g'ri format haqida ma'lumot uchun /help",
                parse_mode=ParseMode.HTML
            )
            os.unlink(tmp_path)
            return
        
        user_id = update.effective_user.id
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_name = f"{user_id}_{timestamp}_{file_name}"
        file_path = upload_dir / unique_name
        shutil.move(tmp_path, str(file_path))
        
        conn = get_db()
        c = conn.cursor()
        
        c.execute('''INSERT INTO files (user_id, file_name, file_path, file_size, original_name, format, question_count)
                     VALUES (?, ?, ?, ?, ?, ?, ?)''',
                  (user_id, unique_name, str(file_path), file_size, file_name, ext[1:].upper(), len(questions)))
        
        file_id = c.lastrowid
        
        for q in questions:
            opts = q['options']
            while len(opts) < 4:
                opts.append(f"Variant {len(opts) + 1}")
            
            c.execute('''INSERT INTO test_questions
                         (file_id, question_text, option_a, option_b, option_c, option_d, correct_answer)
                         VALUES (?, ?, ?, ?, ?, ?, ?)''',
                      (file_id, q['text'][:500], opts[0][:200], opts[1][:200], 
                       opts[2][:200], opts[3][:200], q['correct'][:200]))
        
        conn.commit()
        conn.close()
        
        context.user_data['waiting_for_file'] = False
        
        keyboard = [
            [InlineKeyboardButton("▶️ Test boshlash", callback_data=f"test_{file_id}")],
            [InlineKeyboardButton("🏠 Asosiy menyu", callback_data="back_to_main")]
        ]
        
        await processing_msg.edit_text(
            f"✅ <b>Fayl muvaffaqiyatli yuklandi!</b>\n\n"
            f"📄 Fayl: <code>{file_name}</code>\n"
            f"📊 Savollar: <b>{len(questions)} ta</b>\n"
            f"📁 O'lcham: {file_size // 1024} KB\n\n"
            "Test boshlash uchun tugmani bosing 👇",
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
    
    except Exception as e:
        await processing_msg.edit_text(f"❌ Xatolik: {str(e)[:200]}")
        if 'tmp_path' in locals() and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ==================== FLASK ROUTES ====================

@app.route('/')
def index():
    return jsonify({"status": "ok", "bot": "Test Master Bot"})

@app.route('/health')
def health():
    return jsonify({"status": "healthy"})

@app.route('/webhook', methods=['POST'])
def webhook():
    json_data = request.get_json(force=True)
    if json_data and bot_app:
        update = Update.de_json(json_data, bot_app.bot)
        asyncio.run_coroutine_threadsafe(
            bot_app.process_update(update),
            bot_loop
        )
    return "ok", 200


# ==================== BOT SETUP ====================

bot_app = None
bot_loop = None

def run_bot():
    global bot_app, bot_loop
    
    bot_loop = asyncio.new_event_loop()
    asyncio.set_event_loop(bot_loop)
    
    bot_app = Application.builder().token(TOKEN).build()
    
    bot_app.add_handler(CommandHandler("start", start))
    bot_app.add_handler(CommandHandler("help", help_command))
    bot_app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    bot_app.add_handler(CallbackQueryHandler(button_handler))
    
    async def setup():
        await bot_app.initialize()
        
        webhook_url = WEBHOOK_URL.rstrip('/') + '/webhook'
        await bot_app.bot.delete_webhook(drop_pending_updates=True)
        await bot_app.bot.set_webhook(webhook_url)
        await bot_app.start()
        
        print(f"✅ Webhook o'rnatildi: {webhook_url}")
    
    bot_loop.run_until_complete(setup())
    bot_loop.run_forever()


# ==================== MAIN ====================

if __name__ == '__main__':
    print("🚀 Test Master Bot ishga tushmoqda...")
    
    bot_thread = Thread(target=run_bot, daemon=True)
    bot_thread.start()
    
    print(f"🌐 Web server port {PORT} da ishga tushdi")
    app.run(host='0.0.0.0', port=PORT)
